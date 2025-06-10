from sqlalchemy.orm import Session
from app.db.chroma_client import collection, chroma_client
from app.db.database import get_db
from app.models.document_model import Document
from fastapi import UploadFile
from sentence_transformers import SentenceTransformer
import textwrap
from uuid import uuid4
import pdfplumber
from docx import Document as DocxDocument
import mimetypes

embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

def extract_text_from_pdf(document: UploadFile) -> str:
    """
    Extract text from a PDF document using pdfplumber.

    :param document: UploadFile object representing the PDF document
    :return: Extracted text as a string
    """
    text_content = []
    
    try:
        with pdfplumber.open(document.file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text: 
                    text_content.append(page_text)
        
        return "\n".join(text_content).replace('\x00', '')

    except Exception as e:
        document.file.seek(0) 
        content_bytes = document.file.read()
        try:
            return content_bytes.decode("utf-8").replace('\x00', '')
        except UnicodeDecodeError:
            return content_bytes.decode("latin-1").replace('\x00', '')


def extract_text_from_docx(document: UploadFile) -> str:
    """
    Extract text from a DOCX document.
    
    :param document: UploadFile object representing the DOCX document
    :return: Extracted text as a string
    """

    doc = DocxDocument(document.file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    
    return "\n".join(text)


def extract_text(document: UploadFile) -> str:
    """
    Extract text from a document, attempting to handle various file formats.
    """
    content_type = document.content_type
    filename = document.filename or ""
    
    if content_type in ["application/octet-stream", None]:
        guessed_type, _ = mimetypes.guess_type(filename)
        if guessed_type:
            content_type = guessed_type

    if content_type == "application/pdf":
        return extract_text_from_pdf(document)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(document)
    elif content_type == "text/plain":
        content = document.file.read()
        try:
            return content.decode("utf-8").replace('\x00', '')
        except UnicodeDecodeError:
            return content.decode("latin-1").replace('\x00', '')
    else:
        raise ValueError(f"Unsupported document type: {content_type}")
        

def chunk_text(text: str, chunk_size: int = 512) -> list[str]:
    return textwrap.wrap(text, chunk_size)


def save_document(document: UploadFile, db: Session) -> Document:
    """
    Save a document to the database.
    
    :param db: Database session
    :param document: Document object to save
    :return: Saved document object
    """
    content = extract_text(document)

    uploaded_doc = Document(
        name=document.filename,
        content_type=document.content_type,
        data=content
    )

    db.add(uploaded_doc)
    db.commit()
    db.refresh(uploaded_doc)

    chunks = chunk_text(content)

    embeddings = embedding_model.encode(chunks).tolist()

    chunk_ids = [str(uuid4()) for _ in chunks]
    metadatas = [
        {
            "doc_id": uploaded_doc.id,
            "chunk_index": i,
            "filename": document.filename,
        }
        for i in range(len(chunks))
    ]

    collection.add(
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas,
        ids=chunk_ids,
    )
    
    return uploaded_doc


def get_document(db: Session, document_id: int) -> Document:
    """
    Retrieve a document from the database by its ID.
    
    :param db: Database session
    :param document_id: ID of the document to retrieve
    :return: Document object if found, None otherwise
    """
    return db.query(Document).filter(Document.id == document_id).first()


def delete_document(document_id: int, db: Session) -> bool:
    """
    Delete a document from the database by its ID.
    
    :param db: Database session
    :param document_id: ID of the document to delete
    :return: True if deletion was successful, False otherwise
    """
    document = db.query(Document).filter(Document.id == document_id).first()
    if document:
        results = collection.get(where={"doc_id": document_id})
        ids_to_delete = results.get("ids", [])

        if ids_to_delete:
            collection.delete(ids=ids_to_delete)


        db.delete(document)
        db.commit()
        return True
    return False